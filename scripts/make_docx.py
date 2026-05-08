from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


TITLE = "DR-GSMamba: Distributionally Robust Graph-State-Space Learning for Label-Scarce Hyperspectral Image Classification"
AUTHORS = "Kassim Diarra, Xiaoqiang Di, Boubacar M. Keita, Omar M. Dugsiye and Jinhui Cao"


SECTIONS = [
    (
        "Abstract",
        [
            "Hyperspectral image classification requires models that can identify subtle spectral differences while remaining stable under limited annotations, class imbalance, mixed pixels, and changing training splits. Although convolutional networks, Transformers, graph neural networks, and recent state-space models have improved spectral-spatial representation learning, many methods are still evaluated mainly under fixed splits and provide limited evidence about robustness. This paper proposes DR-GSMamba, a distributionally robust graph-state-space framework for label-scarce hyperspectral image classification. The method integrates a spectral state-space encoder, a compact spatial convolutional stem, patch-level graph message passing, prototype-based classification, and uncertainty-aware robust optimization. The spectral encoder captures long-range band dependencies with linear complexity, while the graph branch models non-local contextual structure inside local neighborhoods. A conditional value-at-risk objective emphasizes difficult samples, and a prototype consistency term improves class compactness and interpretability. The planned evaluation protocol reports overall accuracy, average accuracy, Kappa, Macro-F1, per-class accuracy, uncertainty, computational cost, and multi-seed stability on public HSI benchmarks. The project provides reproducible PyTorch code, automatic table and figure generation, and a complete manuscript pipeline for journal submission.",
            "Index Terms--Hyperspectral image classification, state-space model, Mamba, graph neural network, distributionally robust learning, uncertainty estimation.",
        ],
    ),
    (
        "I. Introduction",
        [
            "Hyperspectral imaging records hundreds of contiguous spectral bands for the same geographic scene. These dense spectral observations make HSI valuable for land-cover mapping, precision agriculture, mineral exploration, urban analysis, and environmental monitoring. In classification tasks, each pixel is assigned to a semantic land-cover category according to its spectral response and spatial context. However, practical HSI classification remains difficult because labeled pixels are expensive to obtain, spectral bands contain redundancy and noise, and several classes share similar reflectance signatures.",
            "Deep learning has become the dominant technical direction for HSI classification. Three-dimensional convolutional neural networks jointly encode spectral and spatial neighborhoods, hybrid CNNs reduce computation by separating spectral and spatial processing, and Transformer-based methods improve global dependency modeling. Graph neural networks further introduce relational reasoning over pixels, superpixels, or regions. More recently, state-space models inspired by Mamba have attracted attention because they can model long spectral sequences with linear complexity. Despite these advances, a persistent weakness remains: many methods report strong accuracy on a fixed split but do not fully analyze stability under scarce labels, random split variation, class imbalance, or hard mixed pixels.",
            "This paper addresses HSI classification from the perspective of robust spectral-spatial learning. The proposed DR-GSMamba framework is designed to improve not only accuracy but also statistical stability and uncertainty awareness. A spectral state-space branch models long-range band dependencies, a spatial convolutional branch extracts local texture, and a graph branch performs message passing between patch nodes. These complementary representations are fused and classified using both a linear head and a prototype uncertainty head. The optimization objective combines cross entropy with a distributionally robust loss that concentrates on hard samples, a prototype consistency term, and graph smoothness regularization.",
            "The main contributions are fourfold. First, a graph-state-space architecture is proposed for label-scarce HSI classification. Second, distributionally robust optimization is introduced to reduce split-dependent instability and improve rare-class behavior. Third, a prototype uncertainty head is designed to expose ambiguous predictions around boundaries and mixed pixels. Fourth, the accompanying codebase supports multi-seed experiments, automatic tables, confusion matrices, and reproducible manuscript assets.",
        ],
    ),
    (
        "II. Related Work",
        [
            "CNN-based HSI classifiers use local convolutional filters to extract spectral-spatial patterns. Models such as 3D CNNs and HybridSN have shown that joint local modeling is effective, but the receptive field is limited and computation can grow quickly with patch size and spectral dimensionality. Residual and dense convolutional designs improve optimization but still depend strongly on local inductive bias.",
            "Transformer-based HSI classifiers use attention to capture long-range dependencies across spectral bands, spatial tokens, or both. Their global modeling ability is attractive, but the quadratic cost of attention and the need for large training sets remain obstacles in label-scarce remote-sensing scenes. Axial, grouped, and lightweight attention variants reduce the cost but do not fully remove the instability problem.",
            "Graph-based methods represent pixels, superpixels, or regions as nodes and use relational message passing to exploit non-local spatial context. These methods are useful for land-cover mapping because neighboring and spectrally similar regions often share semantic structure. However, graph construction can be sensitive to noise and may lack strong spectral sequence modeling.",
            "State-space models offer an efficient alternative for long sequence modeling. For HSI, spectral bands naturally form ordered sequences, making SSM-style encoders suitable for capturing global spectral behavior. The proposed method follows this direction but adds graph reasoning and robust optimization, making the contribution broader than a direct Mamba replacement.",
        ],
    ),
    (
        "III. Proposed Method",
        [
            "Given an HSI cube X with height H, width W, and B spectral bands, the method extracts a center spectrum and a local patch for each labeled pixel. The spectral dimension may be reduced by PCA to suppress redundancy and improve computational efficiency. The model then processes the input through three complementary branches.",
            "The spectral branch treats the center pixel spectrum as an ordered sequence. A lightweight state-space style block applies gated depthwise sequence mixing and residual normalization to capture long-range band dependencies. This branch preserves the sequential nature of the spectrum without using quadratic self-attention.",
            "The spatial branch applies compact two-dimensional convolutions to the local patch. It captures texture, edges, and local neighborhood patterns that are useful for separating land-cover classes with similar spectra. The resulting feature map is also reshaped into patch nodes for graph reasoning.",
            "The graph branch builds a soft adjacency matrix from normalized node similarities and performs message passing within the local patch. This operation allows the classifier to aggregate non-local contextual information inside the patch and reduces sensitivity to isolated noisy pixels.",
            "The fusion module concatenates spectral, spatial, and graph features and projects them into a shared representation. A linear classifier predicts class logits, while a prototype head measures similarity between the fused representation and learned class prototypes. The maximum softmax confidence is used to derive a simple uncertainty score.",
            "The training objective combines cross entropy, conditional value-at-risk robust loss, prototype consistency, and graph smoothness. The robust term focuses learning on high-loss samples, which are often boundary pixels, rare-class samples, mislabeled samples, or spectrally confused examples. This objective is intended to improve both mean accuracy and variance across random training splits.",
        ],
    ),
    (
        "IV. Experimental Setup",
        [
            "The intended evaluation uses Indian Pines, Pavia University, Pavia Centre, Salinas, Houston 2013, and one WHU-Hi dataset when available. For each dataset, the data cube is normalized, reduced by PCA when specified, and divided into train, validation, and test sets by stratified sampling. All experiments should be repeated over multiple random seeds.",
            "The main metrics are overall accuracy, average accuracy, Cohen's Kappa, Macro-F1, per-class accuracy, confusion matrix, mean uncertainty, parameter count, and training/inference time. Reporting Macro-F1 and per-class accuracy is important because HSI datasets are often imbalanced and rare classes can be hidden by high overall accuracy.",
            "Recommended baselines include 3D-CNN, HybridSN, SSRN, SpectralFormer or SSFTT, a graph neural network baseline, a recent Mamba-based HSI classifier, and the previous SCATNet model when appropriate. The ablation study should remove the spectral branch, graph branch, robust loss, prototype head, and graph smoothness term independently.",
            "The current repository includes a synthetic smoke-test configuration to verify the full training and asset-generation pipeline. Final journal results must be produced from real benchmark datasets and should not rely on synthetic results.",
        ],
    ),
    (
        "V. Preliminary Pipeline Verification",
        [
            "A synthetic HSI scene was used only to verify that the implementation runs end-to-end. The smoke test completed training, validation, testing, checkpoint saving, metric logging, LaTeX table generation, and confusion-matrix generation. This confirms the reproducibility pipeline but does not constitute the final experimental evidence for submission.",
            "The synthetic smoke test produced an overall accuracy of 58.75%, average accuracy of 59.81%, Kappa of 53.02%, and Macro-F1 of 58.42% after a short three-epoch run. These values are intentionally not presented as scientific claims; they serve only as execution evidence for the project framework.",
        ],
    ),
    (
        "VI. Conclusion",
        [
            "This paper draft presented DR-GSMamba, a distributionally robust graph-state-space framework for label-scarce hyperspectral image classification. The method combines efficient spectral sequence modeling, local spatial representation learning, graph reasoning, prototype classification, and uncertainty-aware robust optimization. The proposed research direction is designed to support a stronger journal story than a simple accuracy-driven architecture paper by emphasizing stability, rare-class behavior, and reproducibility.",
            "The next step is to run full benchmark experiments, add baseline comparisons, complete statistical significance testing, and replace the preliminary pipeline-verification section with final benchmark results.",
        ],
    ),
]


def set_normal_style(document: Document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)


def main():
    out = Path("paper/DR-GSMamba_manuscript.docx")
    doc = Document()
    set_normal_style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    p = doc.add_paragraph("SUBMISSION TO IEEE TRANSACTIONS ON GEOSCIENCE AND REMOTE SENSING, VOL. XX, NO. XX, XXXX, 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run(AUTHORS)
    run.font.size = Pt(10)

    aff = doc.add_paragraph(
        "The authors are with the School of Computer Science and Technology, Changchun University of Science and Technology, Changchun, China."
    )
    aff.alignment = WD_ALIGN_PARAGRAPH.LEFT
    aff.runs[0].font.size = Pt(8)

    for heading, paragraphs in SECTIONS:
        h = doc.add_heading(heading, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for text in paragraphs:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_section(WD_SECTION.CONTINUOUS)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()

